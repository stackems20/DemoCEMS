from groq import Groq
import pandas as pd
import json
from docx import Document
import re

def answer_generation(data,format,api_key):
    client = Groq(api_key=api_key)
    completion = client.chat.completions.create(
    model="openai/gpt-oss-120b",
    messages=[
        {
            "role": "system",
            "content": """You are an expert Enviromentalist as well as a Process engineer. Your task is to make a reference/helping
            document for CEMS Audit."""
        },
        {
            
            "role": "user",
            "content": f"""Below you are provided with the a format or topics for the Official Audit Report as well as one 
            month CEMS data. Your task is to provide me the text for a document which Audit maker can refer to for their help.

            ##INPUT
            Predicted Data Frame : {data}
            format of Report : {format}

            ##OUTPUT
            Provide only the text which can be directly converted into word document, no extra information needed. Only provide the
            text under heeadings which you can derive from the data and asked in template.
            """
        }
    ],
    temperature=0.01
)

    answer_string = completion.choices[0].message.content
    return answer_string

def add_markdown_table(doc, table_lines):
    """
    Converts markdown table to Word table
    """
    # साफ lines (remove separator like |---|)
    clean_lines = [line for line in table_lines if not re.match(r'^\|\s*-', line)]

    rows = [line.strip().strip('|').split('|') for line in clean_lines]
    rows = [[cell.strip() for cell in row] for row in rows]

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))

    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell


def clean_bold(text):
    """
    Convert **bold** → bold text (basic handling)
    """
    return re.sub(r'\*\*(.*?)\*\*', r'\1', text)


def save_to_word_advanced(text: str, file_name: str = "CEMS_Audit_Report.docx"):
    doc = Document()

    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # -------- TABLE DETECTION --------
        if line.startswith("|") and "|" in line:
            table_lines = []

            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            add_markdown_table(doc, table_lines)
            continue

        # -------- HEADINGS --------
        if line.startswith("###"):
            doc.add_heading(clean_bold(line.replace("###", "").strip()), level=2)

        elif line.startswith("##"):
            doc.add_heading(clean_bold(line.replace("##", "").strip()), level=1)

        elif line.startswith("#"):
            doc.add_heading(clean_bold(line.replace("#", "").strip()), level=0)

        # -------- BULLETS --------
        elif line.startswith("- "):
            doc.add_paragraph(clean_bold(line[2:]), style="List Bullet")

        # -------- NUMBERED LIST --------
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(clean_bold(line), style="List Number")

        # -------- NORMAL TEXT --------
        else:
            doc.add_paragraph(clean_bold(line))

        i += 1

    doc.save(file_name)
    print(f"✅ Document saved as {file_name}")